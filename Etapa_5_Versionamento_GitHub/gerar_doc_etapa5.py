import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def create_etapa5_doc(output_path):
    doc = docx.Document()

    # Configuração de Margens
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Cores do Padrão Senac e do Projeto
    primary_color = RGBColor(0, 51, 102)   # Azul Senac
    secondary_color = RGBColor(235, 130, 20) # Laranja Sorveteria

    # Cabeçalho Principal
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_inst = p_title.add_run("SENAC - TÉCNICO EM DESENVOLVIMENTO DE SISTEMAS\n")
    run_inst.font.name = "Arial"
    run_inst.font.size = Pt(12)
    run_inst.font.bold = True
    run_inst.font.color.rgb = primary_color

    run_title = p_title.add_run("PROJETO INTEGRADOR – ETAPA 5\nGERENCIAMENTO DE CONFIGURAÇÃO E VERSIONAMENTO COM GIT & GITHUB\n")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(15)
    run_title.font.bold = True

    # Caixa de Identificação
    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_info = p_info.add_run(
        "Sistema: Sistema Sorveteria Tropicalzin Cremoso\n"
        "Unidade Curricular: Projeto Integrador Assistente de Desenvolvimento de Sistemas\n"
        "Equipe / Autores: Wallace Teixeira Carvalho e Edcarlos Cardôso de Farias\n"
        "Repositório GitHub: https://github.com/wtc13carvalho-rgb/sorveteria-tropicalzin-cremoso.git\n"
        "Status: Em desenvolvimento (Etapa 5 Concluída)\n"
    )
    r_info.font.size = Pt(10)
    r_info.font.italic = True

    doc.add_paragraph("-" * 80)

    # 1. Contexto e Objetivos
    h1 = doc.add_heading("1. Contexto e Objetivos da Etapa", level=1)
    doc.add_paragraph(
        "Nesta Etapa 5 do Projeto Integrador, o foco central é a proteção do código-fonte contra imprevistos, "
        "o rastreamento histórico de alterações e a implementação de práticas modernas de programação colaborativa em equipe. "
        "Para atingir essas metas, todo o ecossistema desktop desenvolvido nas etapas anteriores (Etapa 1: POO, Etapa 2: UX/UI, "
        "Etapa 3: Telas NetBeans e Etapa 4: Banco de Dados MySQL) foi estruturado e sincronizado em um repositório remoto público no GitHub."
    )

    # 2. Dados do Repositório e Trabalho em Equipe
    doc.add_heading("2. Especificações do Repositório e Colaboração", level=1)
    doc.add_paragraph(
        "O repositório foi configurado para permitir o trabalho colaborativo simultâneo entre os desenvolvedores do time, "
        "com permissões formais de contribuição registradas no GitHub:"
    )

    table_colab = doc.add_table(rows=1, cols=3)
    table_colab.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table_colab.rows[0].cells
    hdr[0].text = "Integrante"
    hdr[1].text = "Usuário GitHub"
    hdr[2].text = "Papel no Repositório"
    for c in hdr:
        set_cell_background(c, "E6F0FA")
        for p in c.paragraphs:
            p.runs[0].font.bold = True

    colaboradores = [
        ("Wallace Teixeira Carvalho", "@wtc13carvalho-rgb", "Criador / Administrador (Owner)"),
        ("Edcarlos Cardôso de Farias", "@edcarloscardoso", "Desenvolvedor / Colaborador Oficial (Write Access)")
    ]
    for nome, user, papel in colaboradores:
        r = table_colab.add_row().cells
        r[0].text = nome
        r[1].text = user
        r[2].text = papel

    # 3. Comandos de Versionamento Aplicados
    doc.add_heading("3. Comandos de Versionamento Aplicados (Git CLI)", level=1)
    doc.add_paragraph(
        "Durante o ciclo de versionamento e sincronização com o GitHub, foram executados os seguintes comandos fundamentais:"
    )

    comandos = [
        ("git init", "Inicialização do repositório local na pasta raiz do projeto integrador."),
        ("git branch -M main", "Definição do branch principal como 'main', atendendo às convenções modernas."),
        ("git remote add origin <URL>", "Vinculação do repositório local ao repositório remoto criado no GitHub."),
        ("git add .", "Rastreamento e indexação de todas as etapas do projeto e do arquivo .gitignore."),
        ("git commit -m '<mensagem>'", "Criação de commits semânticos documentando cada entrega e evolução."),
        ("git push -u origin main", "Envio dos commits e arquivos locais para o servidor remoto do GitHub."),
        ("git pull origin main", "Sincronização de alterações remotas para atualização do ambiente local.")
    ]

    for cmd, desc in comandos:
        p_cmd = doc.add_paragraph()
        r_cmd = p_cmd.add_run(f"• {cmd}: ")
        r_cmd.bold = True
        r_cmd.font.name = "Courier New"
        p_cmd.add_run(desc)

    # 4. Conformidade do Arquivo README.md
    doc.add_heading("4. Conformidade do Arquivo README.md Criado", level=1)
    doc.add_paragraph(
        "Conforme expressamente solicitado no enunciado da atividade, o repositório inclui na sua raiz o arquivo "
        "README.md contemplando rigorosamente os 6 itens avaliativos:"
    )

    itens_readme = [
        ("Título do Projeto", "Sistema Sorveteria Tropicalzin Cremoso 🍨🍦"),
        ("Status do Projeto", "Em desenvolvimento (Etapa 5 – Versionamento com Git & GitHub)"),
        ("Tecnologias Aplicadas", "Java SE, Swing, NetBeans, MySQL 8.0, JDBC e Git/GitHub"),
        ("Time de Desenvolvedores", "Wallace Teixeira Carvalho e Edcarlos Cardôso de Farias"),
        ("Objetivo do Software", "Informatização completa de vendas, estoque de sabores e fluxo de caixa"),
        ("Funcionalidades do Sistema", "Catálogo de requisitos funcionais (RF001 a RF008) e não funcionais")
    ]

    table_readme = doc.add_table(rows=1, cols=2)
    table_readme.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr2 = table_readme.rows[0].cells
    hdr2[0].text = "Requisito da Atividade"
    hdr2[1].text = "Conteúdo Atendido no README.md"
    for c in hdr2:
        set_cell_background(c, "FFF3E6")
        for p in c.paragraphs:
            p.runs[0].font.bold = True

    for req, cont in itens_readme:
        r2 = table_readme.add_row().cells
        r2[0].text = req
        r2[1].text = cont

    # 5. Evidências Visuais (Prints para a Entrega)
    doc.add_heading("5. Evidências de Criação e Uso do Repositório", level=1)
    doc.add_paragraph(
        "Abaixo estão demarcados os espaços para inclusão das capturas de tela (prints) que comprovam a "
        "criação do repositório no GitHub, o histórico de commits e a atuação da equipe:"
    )

    evidencias = [
        ("Evidência 1: Visão Geral do Repositório no GitHub",
         "Captura de tela exibindo a página inicial do repositório remoto no GitHub com o README.md e a árvore de pastas."),
        ("Evidência 2: Histórico de Commits (Git Bash / GitHub)",
         "Captura de tela demonstrando o histórico com mensagens de commit e a autoria dos integrantes da equipe."),
        ("Evidência 3: Convite e Aceite de Colaborador",
         "Captura de tela em Settings > Collaborators comprovando a adição dos membros para desenvolvimento em time.")
    ]

    for titulo_ev, desc_ev in evidencias:
        p_ev = doc.add_paragraph()
        r_ev = p_ev.add_run(f"📷 {titulo_ev}\n")
        r_ev.bold = True
        r_ev.font.color.rgb = primary_color
        p_ev.add_run(f"{desc_ev}\n")

        # Caixa simulada para o print
        tbl_box = doc.add_table(rows=1, cols=1)
        tbl_box.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_box = tbl_box.rows[0].cells[0]
        set_cell_background(cell_box, "F8F9FA")
        p_box = cell_box.paragraphs[0]
        p_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_box = p_box.add_run("\n[ INSERIR PRINT AQUI ]\n")
        r_box.font.italic = True
        r_box.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()

    # 6. Conclusão
    doc.add_heading("6. Conclusão da Etapa 5", level=1)
    doc.add_paragraph(
        "A Etapa 5 foi concluída com êxito, garantindo segurança patrimonial ao código desenvolvido pela equipe e "
        "estabelecendo um fluxo de trabalho profissional baseado em Git. O projeto encontra-se versionado, com documentação "
        "técnica completa e disponível publicamente para avaliação no endereço: "
        "https://github.com/wtc13carvalho-rgb/sorveteria-tropicalzin-cremoso"
    )

    doc.save(output_path)
    print(f"Documentação da Etapa 5 gerada com sucesso em: {output_path}")

if __name__ == "__main__":
    out_dir = "/home/edcarlos/workspace/pessoal/desenvolvimento_atv/Etapa_5_Versionamento_GitHub"
    os.makedirs(out_dir, exist_ok=True)
    out_file = os.path.join(out_dir, "Documentacao_Etapa5_Versionamento_GitHub.docx")
    create_etapa5_doc(out_file)
