import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_etapa1_doc(output_path):
    doc = docx.Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Styles
    primary_color = RGBColor(0, 51, 102) # Azul Senac
    secondary_color = RGBColor(235, 130, 20) # Laranja Sorveteria

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_inst = p_title.add_run("SENAC - TÉCNICO EM DESENVOLVIMENTO DE SISTEMAS\n")
    run_inst.font.name = "Arial"
    run_inst.font.size = Pt(12)
    run_inst.font.bold = True
    run_inst.font.color.rgb = primary_color

    run_title = p_title.add_run("PROJETO INTEGRADOR – ETAPA 1\nDEFINIÇÃO DO SISTEMA E MODELO DE CLASSES ORIENTADO A OBJETOS\n")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(16)
    run_title.font.bold = True

    # Info Box
    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_info = p_info.add_run("Nome do Sistema: Sistema Sorveteria Tropicalzin Cremoso\nResponsável / Equipe: Wallace Teixeira Carvalho / Edcarlos Cardôso de Farias\nVersão: 1.0 | Linguagem: Java SE (Desktop)\n")
    r_info.font.size = Pt(10)
    r_info.font.italic = True

    doc.add_paragraph("-" * 80)

    # Section 1
    h1 = doc.add_heading("1. Contexto e Justificativa do Projeto", level=1)
    p = doc.add_paragraph(
        "O Sistema Sorveteria Tropicalzin Cremoso foi idealizado para atender às necessidades operacionais e "
        "de atendimento de uma sorveteria artesanal em expansão. O mercado atual exige soluções que combinem rapidez "
        "no caixa, controle seguro de sabores e produtos, além de inovação na interação com o cliente. "
        "A empresa enfrenta desafios decorrentes de anotações manuais em comandas de papel, que geram gargalos no balcão, "
        "divergências de estoque e perda de faturamento."
    )

    # Section 2
    doc.add_heading("2. Objetivos do Sistema", level=1)
    p = doc.add_paragraph(
        "Objetivo Geral: Desenvolver um software aplicativo desktop funcional para gestão completa de vendas, estoque "
        "de sabores e atendimento de clientes de uma sorveteria, proporcionando agilidade operacional e controle de dados."
    )
    p_obj = doc.add_paragraph()
    p_obj.add_run("• Automatizar o registro e totalização de pedidos no balcão.\n")
    p_obj.add_run("• Manter o cadastro dinâmico de sabores e produtos em linha.\n")
    p_obj.add_run("• Implementar a funcionalidade inovadora de Sorteio Aleatório de Sabores para o cliente (RF005).\n")
    p_obj.add_run("• Permitir controle de permissões por perfil (Gerente e Atendente).\n")
    p_obj.add_run("• Estabelecer base sólida em Programação Orientada a Objetos (POO) com classes modulares.")

    # Section 3
    doc.add_heading("3. Requisitos do Sistema", level=1)
    doc.add_heading("3.1. Requisitos Funcionais (RF)", level=2)
    
    table_rf = doc.add_table(rows=1, cols=3)
    table_rf.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table_rf.rows[0].cells
    hdr[0].text = "Código"
    hdr[1].text = "Nome do Requisito"
    hdr[2].text = "Descrição Detalhada"

    rfs = [
        ("RF001", "Cadastro de Sabores", "Permitir incluir, editar, inativar e listar sabores (Tradicional, Especial, Zero)."),
        ("RF002", "Cadastro de Produtos", "Gerenciar recipientes/formatos (Casquinha, Copos, Potes, Taças), preços e limites de bolas."),
        ("RF003", "Cadastro de Clientes", "Registrar informações cadastrais de clientes frequentes (Nome, CPF, Telefone, E-mail)."),
        ("RF004", "Registro e Fechamento de Pedidos", "Montar pedido com múltiplos itens, calcular subtotais, descontos e formas de pagamento (PIX, Cartão, Dinheiro)."),
        ("RF005", "Sorteio Aleatório de Sabores", "Funcionalidade interativa para selecionar aleatoriamente N sabores disponíveis no cardápio."),
        ("RF006", "Gerenciamento de Itens do Pedido", "Permitir associar produtos e múltiplos sabores a cada linha do pedido."),
        ("RF007", "Controle de Autenticação", "Tela de login com restrição de permissões entre Atendente e Gerente."),
        ("RF008", "Consulta e Relatório de Vendas", "Exibir listagem de pedidos realizados e métricas de faturamento.")
    ]

    for code, name, desc in rfs:
        row = table_rf.add_row().cells
        row[0].text = code
        row[1].text = name
        row[2].text = desc

    doc.add_heading("3.2. Requisitos Não Funcionais (RNF)", level=2)
    rnfs = [
        ("RNF001", "Plataforma e Arquitetura", "Aplicação desktop desenvolvida em Java (Java SE 17), compatível com NetBeans IDE."),
        ("RNF002", "Usabilidade", "Interface gráfica limpa, intuitiva e ergonômica, otimizada para atendimento rápido em balcão."),
        ("RNF003", "Acessibilidade", "Navegação por teclado (TAB e atalhos), contraste visual em conformidade com diretrizes WCAG."),
        ("RNF004", "Desempenho", "Tempo de resposta imediato para operações em memória e cálculos de valores de pedidos."),
        ("RNF005", "Integridade e Modularidade", "Código estruturado em pacotes e camadas, respeitando encapsulamento e POO.")
    ]
    table_rnf = doc.add_table(rows=1, cols=3)
    table_rnf.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr2 = table_rnf.rows[0].cells
    hdr2[0].text = "Código"
    hdr2[1].text = "Requisito"
    hdr2[2].text = "Especificação"
    for code, name, spec in rnfs:
        row = table_rnf.add_row().cells
        row[0].text = code
        row[1].text = name
        row[2].text = spec

    # Section 4
    doc.add_heading("4. Arquitetura das Classes (POO)", level=1)
    p_arch = doc.add_paragraph(
        "Nesta Etapa 1, foram concebidas e implementadas em Java as classes fundamentais do domínio de negócio, "
        "localizadas no pacote 'com.sorveteria.model'. As classes seguem rigorosamente os princípios de abstração, "
        "encapsulamento (atributos privados com getters e setters) e métodos de conveniência."
    )

    classes_info = [
        ("Usuario.java", "Gerencia os dados de credenciais, nome e perfil (GERENTE/ATENDENTE) para controle de acesso."),
        ("Sabor.java", "Modela os sabores de sorvetes (nome, descrição, tipo e disponibilidade em estoque)."),
        ("Produto.java", "Modela os formatos de venda (Casquinha, Copo, Taça), definindo preço unitário e limite de bolas de sorvete."),
        ("Cliente.java", "Modela o cliente consumidor com nome, CPF, telefone e e-mail."),
        ("ItemPedido.java", "Associa o Produto escolhido, a lista de Sabores selecionados, quantidade e cálculo de subtotal."),
        ("Pedido.java", "Entidade agregadora que gerencia a lista de itens, cliente, atendente, data/hora, forma de pagamento e total."),
        ("SorteadorSabores.java", "Classe utilitária que implementa o algoritmo de embaralhamento e sorteio aleatório de sabores (RF005)."),
        ("TesteModelosEtapa1.java", "Classe executável contendo testes unitários demonstrativos de criação e fluxo das entidades.")
    ]

    for cname, cdesc in classes_info:
        p_c = doc.add_paragraph()
        p_c.add_run(f"• {cname}: ").bold = True
        p_c.add_run(cdesc)

    doc.add_heading("5. Conclusão da Etapa 1", level=1)
    doc.add_paragraph(
        "A Etapa 1 cumpre integralmente os requisitos de concepção do escopo, levantamento de requisitos de software "
        "e estruturação das classes de modelo orientadas a objetos. O código foi compilado e testado com êxito, servindo de base "
        "sólida para o planejamento de UX/UI na Etapa 2 e posterior programação de telas na Etapa 3."
    )

    doc.save(output_path)
    print(f"Documento criado com sucesso em: {output_path}")

if __name__ == "__main__":
    create_etapa1_doc("/home/edcarlos/workspace/pessoal/desenvolvimento_atv/Etapa_1_Definicao_e_Classes/Documentacao_Etapa1_Definicao_e_Classes.docx")
